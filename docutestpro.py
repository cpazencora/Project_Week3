import os
import pandas as pd
import matplotlib.pyplot as plt
from fpdf import FPDF
from docx import Document
from datetime import datetime

# Directory for input and output files
INPUT_DIR = "input"
OUTPUT_DIR = "output"
os.makedirs(INPUT_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Input Layer: Load and normalize CSV
def load_and_normalize_csv(file_name):
    file_path = os.path.join(INPUT_DIR, file_name)
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"CSV file not found in folder: {file_path}")
    df = pd.read_csv(file_path)
    df['status'] = df['status'].str.lower().str.strip()
    df['status'] = df['status'].replace({'pass': 'passed', 'fail': 'failed', 'skip': 'skipped'})
    return df

# Processing Engine
def generate_metrics(df):
    total_tests = len(df)
    passed = len(df[df['status'] == 'passed'])
    failed = len(df[df['status'] == 'failed'])
    skipped = len(df[df['status'] == 'skipped'])
    error_rate = (failed / total_tests) * 100 if total_tests > 0 else 0
    
    metrics = {
        "total_tests": total_tests,
        "passed": passed,
        "failed": failed,
        "skipped": skipped,
        "error_rate": error_rate
    }
    return metrics

# Chart Generator
def create_status_chart(metrics, output_path):
    statuses = ['Passed', 'Failed', 'Skipped']
    counts = [metrics['passed'], metrics['failed'], metrics['skipped']]
    colors = ['green', 'red', 'yellow']  # Colors for Passed, Failed, Skipped

    plt.figure(figsize=(6, 4))
    plt.bar(statuses, counts, color=colors)
    plt.title('Test Status Distribution')
    plt.xlabel('Status')
    plt.ylabel('Number of Tests')
    plt.savefig(output_path)
    plt.close()

# Output Renderer: PDF
def generate_pdf_report(df, metrics, chart_path, output_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    # Summary Section
    pdf.set_font("Arial", style="B", size=14)
    pdf.cell(0, 10, "Test Results Summary", ln=True, align="C")
    pdf.ln(10)
    pdf.set_font("Arial", size=12)

    # Professional summary layout
    summary_lines = [
        f"Total Tests: {metrics['total_tests']}",
        f"Passed: {metrics['passed']} ({metrics['passed'] / metrics['total_tests']:.1%})",
        f"Failed: {metrics['failed']} ({metrics['failed'] / metrics['total_tests']:.1%})",
        f"Skipped: {metrics['skipped']} ({metrics['skipped'] / metrics['total_tests']:.1%})",
        f"Error Rate: {metrics['error_rate']:.2f}%"
    ]
    for line in summary_lines:
        pdf.cell(0, 10, line, ln=True)

    pdf.ln(10)
    if chart_path and os.path.exists(chart_path):
        pdf.image(chart_path, x=60, w=90)

    # Detailed Results Section
    pdf.ln(10)
    pdf.set_font("Arial", style="B", size=14)
    pdf.cell(0, 10, "Detailed Test Results", ln=True)
    pdf.set_font("Arial", size=10)
    pdf.ln(5)

    # Table with professional styling
    col_widths = [60, 30, 30, 60]
    headers = ["Test Case", "Status", "Execution Time", "Comments"]

    # Header row
    pdf.set_fill_color(200, 200, 200)
    for col, width in zip(headers, col_widths):
        pdf.cell(width, 10, col, border=1, align="C", fill=True)
    pdf.ln()

    # Table rows
    pdf.set_font("Arial", size=9)
    for _, row in df.iterrows():
        for col, width in zip(headers, col_widths):
            pdf.cell(width, 10, str(row[col.lower()]), border=1, align="L")
        pdf.ln()

    pdf.output(output_path)

# Output Renderer: Word
def generate_word_report(df, metrics, chart_path, output_path):
    doc = Document()

    # Summary Section
    doc.add_heading("Test Results Summary", level=1)
    summary_lines = [
        ("Total Tests", metrics['total_tests']),
        ("Passed", f"{metrics['passed']} ({metrics['passed'] / metrics['total_tests']:.1%})"),
        ("Failed", f"{metrics['failed']} ({metrics['failed'] / metrics['total_tests']:.1%})"),
        ("Skipped", f"{metrics['skipped']} ({metrics['skipped'] / metrics['total_tests']:.1%})"),
        ("Error Rate", f"{metrics['error_rate']:.2f}%")
    ]
    for key, value in summary_lines:
        doc.add_paragraph(f"{key}: ", style="Normal").add_run(f"{value}").bold = False

    if chart_path and os.path.exists(chart_path):
        doc.add_picture(chart_path, width=doc.sections[0].page_width * 0.5)

    # Detailed Results Section
    doc.add_heading("Detailed Test Results", level=1)
    table = doc.add_table(rows=1, cols=4)

    # Header row with styling
    headers = ["Test Case", "Status", "Execution Time", "Comments"]
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(header)
        run.bold = True

    # Table rows
    for _, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row['test case'])
        cells[1].text = str(row['status'])
        cells[2].text = str(row['execution time'])
        cells[3].text = str(row['comments'])

    # Apply styles to the table
    table.style = 'Table Grid'

    doc.save(output_path)

# Main Function
def main():
    file_name = input(f"Enter the CSV file name (in '{INPUT_DIR}' folder): ")
    df = load_and_normalize_csv(file_name)
    metrics = generate_metrics(df)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    chart_path = os.path.join(OUTPUT_DIR, f"chart_{timestamp}.png")
    pdf_path = os.path.join(OUTPUT_DIR, f"test_report_{timestamp}.pdf")
    word_path = os.path.join(OUTPUT_DIR, f"test_report_{timestamp}.docx")

    # Create chart and embed in reports
    create_status_chart(metrics, chart_path)
    generate_pdf_report(df, metrics, chart_path, pdf_path)
    generate_word_report(df, metrics, chart_path, word_path)

    # Clean up the temporary chart file
    if os.path.exists(chart_path):
        os.remove(chart_path)

    print(f"Reports generated: {pdf_path}, {word_path}")

# Run the Tool
if __name__ == "__main__":
    main()
